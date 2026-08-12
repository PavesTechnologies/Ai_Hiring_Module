import io

import pypdfium2 as pdfium
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.models.candidates import FileFormat
from app.models.jd.job_descriptions import JDSourceFormat


class TextExtractionService:
    """
    Extracts raw text from an uploaded document. Document-type-agnostic —
    the pdfium/python-docx logic has no JD-specific content, so it lives
    under the shared document_processing package (alongside
    StageExecutionService) rather than under app/services/jd/, ready for a
    future Resume pipeline to reuse without relocating it again.
    """

    @staticmethod
    def extract_pdf_text(file_content: bytes) -> str:
        pdf = pdfium.PdfDocument(file_content)
        pages_text = [page.get_textpage().get_text_range() for page in pdf]
        return "\n".join(pages_text)

    @classmethod
    def _iter_docx_block_items(cls, parent):
        """
        Yields Paragraph/Table objects in the order they appear in
        `parent`'s XML body - the standard python-docx recipe for walking
        document order, since Document.paragraphs and Document.tables are
        each a flat, type-only list that silently drops any paragraph
        living inside a table cell. That matters here: a common resume/JD
        layout puts every real section (skills, education, experience) in
        a table for visual columns, and Document.paragraphs alone would
        see only whatever plain text sits outside those tables (typically
        just a name/contact header) - see extract_docx_text.
        """
        # docx.Document is a factory function, not the document class
        # itself (docx.document.Document), so isinstance() can't target it
        # directly - duck-type instead: a Document has `.element.body`, a
        # table Cell has `._tc`.
        parent_elm = parent.element.body if hasattr(parent, "element") else parent._tc
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    @classmethod
    def _docx_block_text(cls, parent) -> list[str]:
        lines = []
        for block in cls._iter_docx_block_items(parent):
            if isinstance(block, Table):
                for row in block.rows:
                    # A horizontally merged cell appears once per grid
                    # column it spans (same underlying <w:tc> element
                    # repeated in row.cells) - track by element identity so
                    # its text isn't duplicated once per spanned column.
                    seen_cells = set()
                    for cell in row.cells:
                        if id(cell._tc) in seen_cells:
                            continue
                        seen_cells.add(id(cell._tc))
                        lines.extend(cls._docx_block_text(cell))
            elif block.text:
                lines.append(block.text)
        return lines

    @classmethod
    def extract_docx_text(cls, file_content: bytes) -> str:
        document = Document(io.BytesIO(file_content))
        return "\n".join(cls._docx_block_text(document))

    @classmethod
    def extract(cls, file_content: bytes, source_format: JDSourceFormat) -> str:
        if source_format == JDSourceFormat.PDF:
            return cls.extract_pdf_text(file_content)
        return cls.extract_docx_text(file_content)

    @classmethod
    def extract_for_resume(cls, file_content: bytes, file_format: FileFormat) -> str:
        """
        FileFormat-dispatched counterpart to extract() for the resume
        pipeline, which validates PDF/DOCX/PNG/JPEG (FileFormat) rather
        than JD's TEXT/PDF/DOCX (JDSourceFormat). PNG/JPEG require OCR,
        which isn't implemented yet — callers should route image-format
        resumes around this method entirely (see
        ResumeProcessingPipeline._mark_ocr_unsupported); this raises rather
        than silently returning empty text if reached directly.
        """
        if file_format == FileFormat.PDF:
            return cls.extract_pdf_text(file_content)
        if file_format == FileFormat.DOCX:
            return cls.extract_docx_text(file_content)
        raise ValueError(
            f"Text extraction for {file_format.value} resumes requires OCR, "
            "which is not yet implemented."
        )

    @staticmethod
    def get_pdf_page_count(file_content: bytes) -> int:
        pdf = pdfium.PdfDocument(file_content)
        return len(pdf)
